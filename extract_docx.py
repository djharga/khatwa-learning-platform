#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
سكريبت لاستخراج النص والصور من ملف Word (.docx) مع الحفاظ على الترتيب
"""

import zipfile
import os
import re
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.oxml import OxmlElement

def get_image_relationships(docx_path):
    """
    الحصول على علاقات الصور من ملف Word
    """
    image_map = {}
    try:
        with zipfile.ZipFile(docx_path, 'r') as zip_ref:
            # قراءة ملف العلاقات
            if 'word/_rels/document.xml.rels' in zip_ref.namelist():
                rels_content = zip_ref.read('word/_rels/document.xml.rels').decode('utf-8')
                # استخراج معرفات الصور
                pattern = r'Id="([^"]+)"[^>]*Target="media/([^"]+)"'
                matches = re.findall(pattern, rels_content)
                for rel_id, filename in matches:
                    image_map[rel_id] = filename
    except Exception as e:
        print(f"خطأ في قراءة العلاقات: {e}")
    return image_map

def extract_images_with_ids(docx_path, output_dir, image_map):
    """
    استخراج الصور مع معرفاتها من ملف Word
    """
    images_dir = Path(output_dir)
    images_dir.mkdir(exist_ok=True)
    
    saved_images = {}
    
    try:
        with zipfile.ZipFile(docx_path, 'r') as zip_ref:
            for rel_id, filename in image_map.items():
                media_path = f'word/media/{filename}'
                if media_path in zip_ref.namelist():
                    ext = os.path.splitext(filename)[1] or '.png'
                    image_path = images_dir / f"image_{rel_id}{ext}"
                    
                    with zip_ref.open(media_path) as source:
                        with open(image_path, 'wb') as target:
                            target.write(source.read())
                    
                    saved_images[rel_id] = {
                        'original': filename,
                        'saved': image_path.name,
                        'path': str(image_path)
                    }
    except Exception as e:
        print(f"خطأ في استخراج الصور: {e}")
    
    return saved_images

def extract_ordered_content(docx_path, images_dir_name, saved_images):
    """
    استخراج المحتوى بالترتيب (نص وصور)
    """
    try:
        doc = Document(docx_path)
        content_items = []
        image_counter = 1
        
        # خريطة لربط معرفات الصور بأسمائها المحفوظة
        rel_id_to_saved = {}
        for rel_id, img_info in saved_images.items():
            rel_id_to_saved[rel_id] = img_info['saved']
        
        for paragraph in doc.paragraphs:
            # التحقق من وجود صور في الفقرة
            paragraph_xml = paragraph._element.xml
            
            # البحث عن صور في الفقرة
            image_ids = re.findall(r'r:embed="([^"]+)"', paragraph_xml)
            image_ids.extend(re.findall(r'r:link="([^"]+)"', paragraph_xml))
            
            # استخراج النص
            para_text = paragraph.text.strip()
            
            # إذا كانت الفقرة تحتوي على صور
            if image_ids:
                for img_id in image_ids:
                    if img_id in rel_id_to_saved:
                        # إضافة الصورة
                        content_items.append({
                            'type': 'image',
                            'content': rel_id_to_saved[img_id],
                            'index': image_counter
                        })
                        image_counter += 1
                    elif img_id in saved_images:
                        # استخدام المعرف المباشر
                        content_items.append({
                            'type': 'image',
                            'content': saved_images[img_id]['saved'],
                            'index': image_counter
                        })
                        image_counter += 1
            
            # إضافة النص إذا كان موجوداً
            if para_text:
                style_name = paragraph.style.name if paragraph.style else ''
                
                # تحديد التنسيق
                if 'Heading' in style_name or 'عنوان' in style_name:
                    level = 1
                    if '1' in style_name:
                        level = 1
                    elif '2' in style_name:
                        level = 2
                    elif '3' in style_name:
                        level = 3
                    content_items.append({
                        'type': 'heading',
                        'content': para_text,
                        'level': level
                    })
                elif paragraph.style.name.startswith('List'):
                    content_items.append({
                        'type': 'list',
                        'content': para_text
                    })
                else:
                    content_items.append({
                        'type': 'text',
                        'content': para_text
                    })
        
        # استخراج الصور من الجداول أيضاً
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_xml = cell._element.xml
                    image_ids = re.findall(r'r:embed="([^"]+)"', cell_xml)
                    for img_id in image_ids:
                        if img_id in rel_id_to_saved:
                            content_items.append({
                                'type': 'image',
                                'content': rel_id_to_saved[img_id],
                                'index': image_counter
                            })
                            image_counter += 1
                    
                    cell_text = cell.text.strip()
                    if cell_text:
                        content_items.append({
                            'type': 'text',
                            'content': cell_text
                        })
        
        return content_items
    
    except Exception as e:
        print(f"خطأ في استخراج المحتوى: {e}")
        import traceback
        traceback.print_exc()
        return []

def create_organized_markdown(content_items, images_dir_name):
    """
    إنشاء ملف Markdown منظم مع ربط النص بالصور
    """
    markdown_content = "# تعديلات الموقع\n\n"
    
    current_section = None
    image_buffer = []
    text_buffer = []
    
    for item in content_items:
        if item['type'] == 'image':
            # إذا كانت هناك نصوص مكدسة، أضفها أولاً
            if text_buffer:
                markdown_content += '\n'.join(text_buffer) + '\n\n'
                text_buffer = []
            
            # إضافة الصورة
            markdown_content += f"![صورة {item['index']}]({images_dir_name}/{item['content']})\n\n"
            image_buffer.append(item['index'])
        
        elif item['type'] == 'heading':
            # حفظ العنوان الحالي
            if text_buffer:
                markdown_content += '\n'.join(text_buffer) + '\n\n'
                text_buffer = []
            
            level = item.get('level', 1)
            markdown_content += f"{'#' * level} {item['content']}\n\n"
            current_section = item['content']
        
        elif item['type'] == 'list':
            text_buffer.append(f"- {item['content']}")
        
        elif item['type'] == 'text':
            # إضافة النص بعد الصورة مباشرة
            text_buffer.append(item['content'])
    
    # إضافة أي نصوص متبقية
    if text_buffer:
        markdown_content += '\n'.join(text_buffer) + '\n\n'
    
    return markdown_content

def main():
    docx_file = "1212.docx"
    output_text_file = "1212.md"
    images_dir = "1212_images"
    
    if not os.path.exists(docx_file):
        print(f"الملف {docx_file} غير موجود!")
        return
    
    print(f"جارٍ معالجة الملف: {docx_file}")
    
    # الحصول على علاقات الصور
    print("جارٍ قراءة علاقات الصور...")
    image_map = get_image_relationships(docx_file)
    print(f"تم العثور على {len(image_map)} صورة في العلاقات")
    
    # استخراج الصور
    print("جارٍ استخراج الصور...")
    saved_images = extract_images_with_ids(docx_file, images_dir, image_map)
    print(f"تم استخراج {len(saved_images)} صورة")
    
    # استخراج المحتوى بالترتيب
    print("جارٍ استخراج المحتوى بالترتيب...")
    content_items = extract_ordered_content(docx_file, images_dir, saved_images)
    print(f"تم استخراج {len(content_items)} عنصر")
    
    # إنشاء ملف Markdown منظم
    print("جارٍ إنشاء ملف Markdown منظم...")
    markdown_content = create_organized_markdown(content_items, images_dir)
    
    # حفظ الملف
    with open(output_text_file, 'w', encoding='utf-8') as f:
        f.write(markdown_content)
    
    print(f"\n✅ تم الانتهاء بنجاح!")
    print(f"📄 ملف النص: {output_text_file}")
    print(f"🖼️  مجلد الصور: {images_dir}/")
    print(f"📊 عدد الصور: {len(saved_images)}")
    print(f"📝 عدد العناصر: {len(content_items)}")

if __name__ == "__main__":
    main()
